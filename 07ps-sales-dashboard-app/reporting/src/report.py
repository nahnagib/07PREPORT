"""Builds the management-ready .docx: title page, executive summary, forecast
charts with confidence bands, key observations, driver analysis table, risks,
and an appendix of raw KPI tables.
"""
from __future__ import annotations

import datetime as dt
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from . import drivers as drivers_mod
from . import forecast as forecast_mod
from .stats import Scorecard


def _chart_forecast(daily: pd.DataFrame, fc: forecast_mod.ForecastResult, title: str, out_path: Path) -> Path | None:
    if fc.method == "none" or not fc.point_forecast:
        return None
    fig, ax = plt.subplots(figsize=(8, 4))
    if not daily.empty:
        history = daily.tail(120)
        ax.plot(history["date"], history["value"], label="Actual", color="#2C6E9E")
    ax.plot(fc.horizon_dates, fc.point_forecast, label=f"Forecast ({fc.method})", color="#D9822B")
    ax.fill_between(fc.horizon_dates, fc.lower_ci, fc.upper_ci, color="#D9822B", alpha=0.2, label="80% confidence")
    ax.set_title(title)
    ax.legend(loc="upper left", fontsize=8)
    ax.tick_params(axis="x", rotation=30)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)
    return out_path


def _chart_monthly(panel: pd.DataFrame, fc: forecast_mod.ForecastResult, col: str, title: str, out_path: Path) -> Path | None:
    if panel.empty:
        return None
    fig, ax = plt.subplots(figsize=(8, 4))
    dates_hist = pd.to_datetime(panel["year_month"] + "-01")
    ax.plot(dates_hist, panel[col], label="Actual", color="#2C6E9E", marker="o", markersize=3)
    if fc.method != "none" and fc.point_forecast:
        ax.plot(fc.horizon_dates, fc.point_forecast, label=f"Forecast ({fc.method})", color="#D9822B", marker="o", markersize=3)
        ax.fill_between(fc.horizon_dates, fc.lower_ci, fc.upper_ci, color="#D9822B", alpha=0.2, label="80% confidence")
    ax.set_title(title)
    ax.legend(loc="upper left", fontsize=8)
    ax.tick_params(axis="x", rotation=30)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)
    return out_path


def build_report(
    output_path: Path,
    chart_dir: Path,
    anchor_date: dt.date,
    sc: Scorecard,
    revenue_forecast: forecast_mod.ForecastResult,
    pipeline_forecast: forecast_mod.ForecastResult,
    monthly_panel: pd.DataFrame,
    exec_summary: str,
    observations: list[str],
    outlook_lines: list[str],
    risks: list[str],
    gap_contributions: list[drivers_mod.GapContribution],
    shap_findings: list,
    shap_note: str,
    bottlenecks: list[drivers_mod.StageBottleneck],
    methodology: str,
) -> None:
    chart_dir.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # ---- Title page ----
    title = doc.add_heading("Sales Predictive Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"Status as of {anchor_date.isoformat()}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    doc.add_page_break()

    # ---- Executive Summary ----
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(exec_summary)

    # ---- Forecast charts ----
    doc.add_heading("Forecast Charts", level=1)
    revenue_chart = _chart_forecast(sc.daily_revenue, revenue_forecast, "Daily Revenue: Actual & 90-Day Forecast", chart_dir / "revenue_forecast.png")
    if revenue_chart:
        doc.add_picture(str(revenue_chart), width=Inches(6.5))
    else:
        doc.add_paragraph(f"Revenue forecast chart unavailable: {revenue_forecast.note}")

    pipeline_chart = _chart_monthly(monthly_panel, pipeline_forecast, "opportunities", "Monthly Opportunities: Actual & Forecast", chart_dir / "pipeline_forecast.png")
    if pipeline_chart:
        doc.add_picture(str(pipeline_chart), width=Inches(6.5))

    # ---- Key Observations ----
    doc.add_heading("Key Observations", level=1)
    for obs in observations:
        doc.add_paragraph(obs, style="List Bullet")

    # ---- Forecast Outlook ----
    doc.add_heading("Forecast Outlook", level=1)
    for line in outlook_lines:
        doc.add_paragraph(line, style="List Bullet")

    # ---- Risks / Watch List ----
    doc.add_heading("Risks / Areas Needing Attention", level=1)
    for r in risks:
        doc.add_paragraph(r, style="List Bullet")

    # ---- Driver Analysis ----
    doc.add_heading("Driver Analysis", level=1)
    doc.add_heading("Revenue Gap by Salesperson (YTD, worst first)", level=2)
    if gap_contributions:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Salesperson", "Actual", "Target", "Gap"
        for g in gap_contributions[:15]:
            row = table.add_row().cells
            row[0].text = g.salesperson
            row[1].text = f"${g.actual:,.0f}"
            row[2].text = f"${g.target:,.0f}"
            row[3].text = f"${g.gap:,.0f} ({g.pct_of_total_gap:.0f}% of shortfall)" if g.gap < 0 else f"${g.gap:,.0f}"
    else:
        doc.add_paragraph("No salesperson-level target data available for this period.")

    doc.add_heading("Lost-Opportunity Drivers", level=2)
    doc.add_paragraph(shap_note)
    if shap_findings:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Factor", "Value", "Impact"
        for f in shap_findings:
            row = table.add_row().cells
            row[0].text = f.factor.replace("_", " ").title()
            row[1].text = f.value_label
            row[2].text = f"{f.impact_pct:.1f}%"

    doc.add_heading("Pipeline Stage Volumes", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Transition", "In", "Out", "Ratio"
    for b in bottlenecks:
        row = table.add_row().cells
        row[0].text = b.transition
        row[1].text = str(b.actual_count_in)
        row[2].text = str(b.actual_count_out)
        row[3].text = f"{b.conversion_pct:.0f}%"

    # ---- Appendix ----
    doc.add_page_break()
    doc.add_heading("Appendix: KPI Scorecard", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = "KPI", "Actual", "Target", "vs LY", "Status"
    for card in sc.cards.values():
        row = table.add_row().cells
        row[0].text = card.name
        row[1].text = f"${card.actual:,.0f}" if card.unit == "currency" else f"{card.actual:,.1f}"
        row[2].text = (f"${card.target:,.0f}" if card.unit == "currency" else f"{card.target:,.1f}") if card.target else "--"
        row[3].text = f"{card.variance_vs_last_year_pct:+.1f}%" if card.variance_vs_last_year_pct is not None else "--"
        row[4].text = card.status.replace("_", " ").title()

    if sc.unavailable:
        doc.add_heading("Unavailable Sections", level=2)
        for section, msg in sc.unavailable.items():
            doc.add_paragraph(f"{section}: {msg}", style="List Bullet")

    doc.add_heading("Methodology", level=2)
    doc.add_paragraph(methodology)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
